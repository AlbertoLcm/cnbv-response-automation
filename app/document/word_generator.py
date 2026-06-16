"""
word_generator.py
=================
Generación de documentos Word (.docx) con el formato SITI/BBVA a partir de
un diccionario de variables extraídas del JSON de datos.

Funciones de formato:
    - set_col_widths              → Fija anchos de columnas en una tabla.
    - set_row_heights_and_align   → Fija altura de filas y alineación vertical.
    - set_cell_background         → Aplica color de fondo a una celda.

Función principal:
    - generar_documento_siti      → Construye el documento completo y lo guarda.
"""

import os
from io import BytesIO
from datetime import datetime

import requests
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# URL del logo corporativo
# ---------------------------------------------------------------------------
_LOGO_URL = "https://albertolcm.github.io/public-images/imagenes/bbva.png"

# Datos fijos del bloque de firma
_DATOS_FIRMA = [
    "VANESSA RAMIREZ PATIÑO",
    "GERENTE",
    "BBVA MÉXICO, S.A.",
    "INSTITUCIÓN DE BANCA MÚLTIPLE",
    "GRUPO FINANCIERO BBVA MÉXICO",
]

# Mapeos de localización para fechas en español
_DIAS = {0: "Lunes", 1: "Martes", 2: "Miércoles", 3: "Jueves",
         4: "Viernes", 5: "Sábado", 6: "Domingo"}
_MESES = {1: "Enero", 2: "Febrero", 3: "Marzo", 4: "Abril",
          5: "Mayo", 6: "Junio", 7: "Julio", 8: "Agosto",
          9: "Septiembre", 10: "Octubre", 11: "Noviembre", 12: "Diciembre"}


# ---------------------------------------------------------------------------
# Helpers de formato de tablas
# ---------------------------------------------------------------------------

def set_col_widths(table, widths: list) -> None:
    """Aplica los anchos de columna indicados a todas las filas de la tabla."""
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def set_row_heights_and_align(table, height) -> None:
    """
    Establece la altura de cada fila y la alineación vertical de sus celdas
    en BOTTOM (alineado al pie).
    """
    for row in table.rows:
        row.height = height
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM


def set_cell_background(cell, color_hex: str) -> None:
    """
    Aplica un color de fondo (hex sin '#') a la celda indicada.

    Args:
        cell:      Celda de una tabla python-docx.
        color_hex: Color en formato hexadecimal sin '#', ej. 'A6A6A6'.
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    shd   = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  color_hex)
    tc_pr.append(shd)


def _add_cell_borders(cell, border_names: list[str] = None) -> None:
    """Añade bordes simples a los lados indicados de una celda."""
    if border_names is None:
        border_names = ["top", "left", "bottom", "right"]
    tc_pr     = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for name in border_names:
        tag = OxmlElement(f"w:{name}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"),  "4")
        tc_borders.append(tag)
    tc_pr.append(tc_borders)


def _remove_cell_borders(cell) -> None:
    """Elimina todos los bordes de una celda."""
    tc_pr     = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for name in ["top", "left", "bottom", "right"]:
        border_elm = OxmlElement(f"w:{name}")
        border_elm.set(qn("w:val"), "none")
        tc_borders.append(border_elm)
    tc_pr.append(tc_borders)


# ---------------------------------------------------------------------------
# Generador del documento
# ---------------------------------------------------------------------------

def generar_documento_siti(variables: dict, dir_salida: str) -> str | None:
    """
    Genera un documento Word (.docx) con el formato SITI/BBVA usando los
    datos proporcionados en el diccionario `variables`.

    Args:
        variables:  Diccionario con los datos del oficio. Claves esperadas:
                    oficio, expediente, folio, expediente_corto, autoridad,
                    tipo_respuesta, tipo_asunto, nombre, rfc, variable_dof,
                    fecha (opcional, usa fecha actual si no se provee),
                    no_cuenta, tipo_cuenta, estatus_cuenta, caracter_cuenta,
                    sucursal, saldo, moneda, observaciones.
        dir_salida: Directorio donde se guardará el archivo .docx generado.

    Returns:
        Ruta completa del archivo generado, o None si ocurre un error al guardar.
    """
    doc = Document()

    # --- Estilo base ---
    estilo = doc.styles["Normal"]
    estilo.font.name = "Helvetica"
    estilo.font.size = Pt(9)

    # --- Márgenes de página ---
    section = doc.sections[0]
    section.top_margin    = Cm(3.53)
    section.bottom_margin = Cm(1.76)
    section.left_margin   = Cm(1.41)
    section.right_margin  = Cm(1.41)

    # =========================================================
    # 1. LOGO EN EL ENCABEZADO
    # =========================================================
    try:
        resp = requests.get(
            _LOGO_URL,
            headers={"User-Agent": "Mozilla/5.0"},
            timeout=15,
        )
        if resp.status_code == 200:
            img_data = BytesIO(resp.content)
            img_data.seek(0)

            header     = section.header
            hdr_table  = header.add_table(1, 1, Inches(5))
            hdr_table.autofit = False
            hdr_cell   = hdr_table.cell(0, 0)

            _remove_cell_borders(hdr_cell)

            p_hdr = hdr_cell.paragraphs[0]
            p_hdr.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_hdr.add_run().add_picture(img_data, width=Inches(1.5))
            p_hdr.paragraph_format.space_after = Pt(0)
        else:
            print(f"[word_generator] Error HTTP {resp.status_code} al descargar el logo.")
            doc.add_paragraph(f"<<< ERROR HTTP {resp.status_code} AL OBTENER LOGO >>>")
    except Exception as exc:
        print(f"[word_generator] Error de red al obtener logo: {exc}")
        doc.add_paragraph(f"<<< ERROR AL OBTENER LOGO: {exc} >>>")

    doc.add_paragraph()

    # =========================================================
    # 2. FECHA
    # =========================================================
    ahora     = datetime.now()
    fecha_hoy = f"{_DIAS[ahora.weekday()]} {ahora.day} de {_MESES[ahora.month]} de {ahora.year}"

    p_fecha = doc.add_paragraph(variables.get("fecha", fecha_hoy))
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()

    # =========================================================
    # 3. ENCABEZADO + RECUADRO ACUSE
    # =========================================================
    tbl_enc = doc.add_table(rows=1, cols=2)
    tbl_enc.autofit = False
    tbl_enc.rows[0].height = Cm(1.8)
    set_col_widths(tbl_enc, [Cm(11.48), Cm(7.252)])

    # Celda izquierda: datos institucionales
    celda_izq = tbl_enc.cell(0, 0)
    p_izq = celda_izq.paragraphs[0]
    run_izq = p_izq.add_run(
        "Comisión Nacional Bancaria y de Valores\n"
        "Vicepresidencia de Supervisión de Procesos Preventivos\n"
        "Dirección General de Atención a Autoridades\n"
        'Coordinación de Atención a Autoridades "A"'
    )
    run_izq.bold = True
    run_izq.font.size = Pt(10)

    # Celda derecha: acuse de recibo con borde
    celda_der = tbl_enc.cell(0, 1)
    celda_der.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_der = celda_der.paragraphs[0]
    p_der.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_acuse = p_der.add_run("\nAcuse de recibo CNBV")
    run_acuse.bold = True
    run_acuse.font.color.rgb = RGBColor(166, 166, 166)
    _add_cell_borders(celda_der)

    doc.add_paragraph()

    p_atencion = doc.add_paragraph("Atención Dr. Jorge Alfredo Ramírez Talamantes")
    p_atencion.runs[0].bold = True

    # =========================================================
    # 4. TABLA DE DATOS DEL OFICIO
    # =========================================================
    tbl_oficio = doc.add_table(rows=4, cols=3)
    tbl_oficio.style   = "Table Grid"
    tbl_oficio.autofit = False
    set_col_widths(tbl_oficio, [Cm(2.0), Cm(4.5), Cm(12.25)])
    set_row_heights_and_align(tbl_oficio, Cm(0.6))

    # Columna "Asunto" abarca las 4 filas
    celda_asunto = tbl_oficio.cell(0, 0)
    celda_asunto.merge(tbl_oficio.cell(3, 0))
    celda_asunto.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    celda_asunto.text = "Asunto:"
    p_asunto = celda_asunto.paragraphs[0]
    p_asunto.runs[0].bold = True
    p_asunto.alignment    = WD_ALIGN_PARAGRAPH.CENTER

    filas_oficio = [
        ("Oficio:",               variables.get("oficio",     "")),
        ("Expediente:",           variables.get("expediente", "")),
        ("Folio:",                variables.get("folio",      "")),
        ("Autoridad solicitante:", variables.get("autoridad",  "")),
    ]
    for fila_idx, (label, valor) in enumerate(filas_oficio):
        tbl_oficio.cell(fila_idx, 1).text = label
        tbl_oficio.cell(fila_idx, 2).text = valor

    doc.add_paragraph()

    # =========================================================
    # 5. TABLA DE TIPOS
    # =========================================================
    tbl_tipos = doc.add_table(rows=2, cols=2)
    tbl_tipos.style   = "Table Grid"
    tbl_tipos.autofit = False
    set_col_widths(tbl_tipos, [Cm(4.5), Cm(14.25)])
    set_row_heights_and_align(tbl_tipos, Cm(0.6))

    for fila_idx, (label, key) in enumerate(
        [("Tipo de respuesta:", "tipo_respuesta"), ("Tipo de asunto:", "tipo_asunto")]
    ):
        celda_label = tbl_tipos.cell(fila_idx, 0)
        celda_label.text = label
        celda_label.paragraphs[0].runs[0].bold = True
        tbl_tipos.cell(fila_idx, 1).text = variables.get(key, "")

    doc.add_paragraph()

    doc.add_paragraph(
        "En atención al oficio señalado al rubro, nos permitimos hacer de su conocimiento "
        "que se ha procedido a ejecutar la instrucción de la autoridad requirente respecto "
        "a la(s) persona(s) que abajo se indica(n):"
    )

    # =========================================================
    # 6. TABLA NOMBRE / RFC
    # =========================================================
    tbl_cliente = doc.add_table(rows=1, cols=4)
    tbl_cliente.style   = "Table Grid"
    tbl_cliente.autofit = False
    set_col_widths(tbl_cliente, [Cm(2.5), Cm(9.5), Cm(1.5), Cm(5.25)])
    set_row_heights_and_align(tbl_cliente, Cm(0.6))

    datos_cliente = [
        ("NOMBRE:", variables.get("nombre", "")),
        ("RFC:",    variables.get("rfc",    "")),
    ]
    col_map = [0, 1, 2, 3]
    tbl_cliente.cell(0, 0).text = "NOMBRE:"
    tbl_cliente.cell(0, 1).text = variables.get("nombre", "")
    tbl_cliente.cell(0, 2).text = "RFC:"
    tbl_cliente.cell(0, 3).text = variables.get("rfc", "")

    color_gris = "A6A6A6"
    for row in tbl_cliente.rows:
        for cell in row.cells:
            set_cell_background(cell, color_gris)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

    doc.add_paragraph()

    # =========================================================
    # 7. TABLA DE CUENTAS
    # =========================================================
    tbl_cuentas = doc.add_table(rows=3, cols=7)
    tbl_cuentas.style   = "Table Grid"
    tbl_cuentas.autofit = False
    set_col_widths(
        tbl_cuentas,
        [Cm(2.5), Cm(3.2), Cm(1.8), Cm(1.8), Cm(4.7), Cm(2.5), Cm(2.25)],
    )

    encabezados = [
        "No. Cuenta", "Tipo", "Estatus", "Carácter",
        "Ubicación/Sucursal", "Saldo", "Moneda",
    ]
    for i, texto in enumerate(encabezados):
        celda = tbl_cuentas.cell(0, i)
        celda.text = texto
        p = celda.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        set_cell_background(celda, "D9D9D9")

    datos_cuenta = [
        variables.get("no_cuenta",       ""),
        variables.get("tipo_cuenta",     ""),
        variables.get("estatus_cuenta",  ""),
        variables.get("caracter_cuenta", ""),
        variables.get("sucursal",        ""),
        variables.get("saldo",           ""),
        variables.get("moneda",          ""),
    ]
    for i, texto in enumerate(datos_cuenta):
        celda = tbl_cuentas.cell(1, i)
        celda.text = texto
        celda.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Fila de observaciones
    celda_obs_titulo = tbl_cuentas.cell(2, 0)
    celda_obs_titulo.text = "Observaciones"
    celda_obs_titulo.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    celda_obs_texto = tbl_cuentas.cell(2, 1)
    celda_obs_texto.merge(tbl_cuentas.cell(2, 6))
    celda_obs_texto.text = variables.get("observaciones", "")
    celda_obs_texto.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    for row in tbl_cuentas.rows:
        row.height = Cm(0.8)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    # =========================================================
    # 8. BLOQUE DE FIRMA
    # =========================================================
    p_linea = doc.add_paragraph("______________________________________")
    p_linea.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for linea in _DATOS_FIRMA:
        p_data = doc.add_paragraph(linea)
        p_data.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_data.paragraph_format.space_before = Pt(0)
        p_data.paragraph_format.space_after  = Pt(0)

    doc.add_paragraph()

    p_dof = doc.add_paragraph(f"{variables.get('variable_dof', '')}-AJRT")
    p_dof.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # =========================================================
    # 9. GUARDAR
    # =========================================================
    try:
        os.makedirs(dir_salida, exist_ok=True)
        expediente_corto = variables.get("expediente_corto", "SIN_EXPEDIENTE")
        nombre_archivo   = f"{expediente_corto}.docx"
        ruta_completa    = os.path.join(dir_salida, nombre_archivo)

        doc.save(ruta_completa)
        print(f"[word_generator] ✅ Documento generado: {ruta_completa}")
        return ruta_completa

    except Exception as exc:
        print(f"[word_generator] ❌ Error al guardar el documento: {exc}")
        return None
