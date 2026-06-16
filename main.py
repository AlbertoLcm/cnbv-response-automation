import os
import shutil
import re 
import docx
import pandas as pd
import xml.etree.ElementTree as ET
import json
import sys

#################################
from playwright.async_api import (
    async_playwright,
    Page,
    TimeoutError as PlaywrightTimeoutError,
)
from tqdm.asyncio import tqdm
import asyncio
import getpass


from io import BytesIO
import requests
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime


from playwright.sync_api import sync_playwright, TimeoutError as PlaywrightTimeoutError


FOLDER_GENERATES = 'dist'

FOLDER_DOCUMENTS = os.path.join(FOLDER_GENERATES, 'documents')
INPUT_FILE = 'layout.xlsx'
COLUMS_REQUIRED = ["Oficio", "Expediente", "Publicacion", "Plazo", "Expediente Corto", "Caratula", "XML", "Archivo Doc", "Archivo XML"]
COLUMNS_WITH_SUGO = ["Oficio", "Folio Sugo", "Extract Folio Sugo", "Expediente", "Publicacion", "Plazo", "Expediente Corto", "Caratula", "XML", "Archivo Doc", "Archivo XML"]
OUTPUT_EXCEL = os.path.join(FOLDER_GENERATES, 'layout_folio_sugo.xlsx')
OUTPUT_JSON = os.path.join(FOLDER_GENERATES, 'extract_data.json')
OUTPUT_TEMPORAL = os.path.join(FOLDER_GENERATES, "resultados_folio_sugo_temp.csv")

BATCH_GUARDADO = 25
CHECKPOINT_EVERY = 25
NUM_PESTANAS = 5

URL_LOGIN = "https://acprod.intranet.com.mx/mbom_mx_ws/mbom_mx_web/PortalLogon"
URL_ESTATUS_FOLIO = "https://acprod.intranet.com.mx:443/boixp_mx_web/boixp_mx_web/servlet/ServletOperacionWeb?OPERACION=VGOMX012&LOCALE=es_ES&DATOS_ENTRADA.FLUJO_LANZAR=GOMXFL10090"
URL_CONSULTA_CEDULA = "https://acprod.intranet.com.mx/boixp_mx_web/boixp_mx_web/servlet/ServletOperacionWeb?OPERACION=VGOMX042&LOCALE=es_ES&DATOS_ENTRADA.FLUJO_LANZAR=GOMXFL10190"
URL_SUGO = "https://acprod.intranet.com.mx/mbom_mx_ws/mbom_mx_web/mbom_mx_web_jsp/portal3.jsp"


# --- CONFIGURACIÓN DOWNLOAD FILES CNBV ---
ARCHIVO_CREDENCIALES = 'credenciales.json'
ARCHIVO_RESULTADOS = 'layout_descargas.xlsx'

# --- UTILIDADES ---
def split_list(lista, n):
    k, m = divmod(len(lista), n)
    return (lista[i * k + min(i, m) : (i + 1) * k + min(i + 1, m)] for i in range(n))

def print_banner_ascii():
    banner = """
    ██╗███╗   ██╗███████╗ ██████╗ ██████╗ ███╗   ███╗███████╗███████╗
    ██║████╗  ██║██╔════╝██╔═══██╗██╔══██╗████╗ ████║██╔════╝██╔════╝
    ██║██╔██╗ ██║█████╗  ██║   ██║██████╔╝██╔████╔██║█████╗  ███████╗
    ██║██║╚██╗██║██╔══╝  ██║   ██║██╔══██╗██║╚██╔╝██║██╔══╝  ╚════██║
    ██║██║ ╚████║██║     ╚██████╔╝██║  ██║██║ ╚═╝ ██║███████╗███████║
    ╚═╝╚═╝  ╚═══╝╚═╝      ╚═════╝ ╚═╝  ╚═╝╚═╝     ╚═╝╚══════╝╚══════╝
                            v1.0.0 | Generador de Informes
    """
    print(banner)
    print("-" * 65)

"""
    ==================================
    FUNCIONES EXTRACTOR DATA
    ==================================
"""

def extractor_text_document_word(document_path):
    """
        Extrae el texto completo del documento bajando al XML subyacente.
        Esto permite leer Controles de Contenido, Campos de Formulario y Textboxes
        que la función normal de python-docx ignora.
    """
    try:
        doc = docx.Document(document_path)
        text_document = []
        
        for element in doc.element.body.iter():
            if element.tag.endswith('}t'):
                if element.text:
                    text_document.append(element.text)
            elif element.tag.endswith('}p'):
                text_document.append("\n")
                
        return "".join(text_document)
    except Exception as e:
        print(f"Error al leer {document_path}: {e}")
        return ""
    
def search_autoridad(text):
    """ 
        Extrae el nombre completo de la autoridad a partir de un texto.
    """
    patron = r"emitido por\s+(.+?),\s*a efecto"
    coincidencia = re.search(patron, text, re.IGNORECASE)
    
    if coincidencia:
        autoridad_extraida = coincidencia.group(1).strip()
        return autoridad_extraida
    return "No encontrado"

def generar_estructura_vacia_xml():
    """
        Genera la misma estructura base del XML pero con datos vacíos.
        Se utiliza cuando no hay XML asociado o el archivo no existe.
    """
    return {
        "DatosGenerales": {
            "Cnbv_NumeroOficio": "",
            "Cnbv_NumeroExpediente": "",
            "Cnbv_SolicitudSiara": "",
            "Cnbv_Folio": "",
            "Cnbv_OficioYear": "",
            "Cnbv_AreaClave": "",
            "Cnbv_AreaDescripcion": "",
            "Cnbv_FechaPublicacion": "",
            "Cnbv_DiasPlazo": "",
            "AutoridadNombre": "",
            "NombreSolicitante": "",
            "Referencia": "",
            "Referencia1": "",
            "Referencia2": "",
            "TieneAseguramiento": ""
        },
        "SolicitudPartes": [],
        "SolicitudesEspecificas": []
    }

def parsear_expediente_cnbv(xml_source, es_archivo=False):
    if es_archivo:
        tree = ET.parse(xml_source)
        root = tree.getroot()
    else:
        root = ET.fromstring(xml_source)
    
    ns = {'ns': 'http://www.cnbv.gob.mx'}
    
    def obtener_texto_limpio(elemento, tag):
        nodo = elemento.find(f'ns:{tag}', ns)
        if nodo is not None and nodo.text is not None:
            return nodo.text.strip()
        return ""

    datos_resultado = {
        "DatosGenerales": {
            "Cnbv_NumeroOficio": obtener_texto_limpio(root, "Cnbv_NumeroOficio"),
            "Cnbv_NumeroExpediente": obtener_texto_limpio(root, "Cnbv_NumeroExpediente"),
            "Cnbv_SolicitudSiara": obtener_texto_limpio(root, "Cnbv_SolicitudSiara"),
            "Cnbv_Folio": obtener_texto_limpio(root, "Cnbv_Folio"),
            "Cnbv_OficioYear": obtener_texto_limpio(root, "Cnbv_OficioYear"),
            "Cnbv_AreaClave": obtener_texto_limpio(root, "Cnbv_AreaClave"),
            "Cnbv_AreaDescripcion": obtener_texto_limpio(root, "Cnbv_AreaDescripcion"),
            "Cnbv_FechaPublicacion": obtener_texto_limpio(root, "Cnbv_FechaPublicacion"),
            "Cnbv_DiasPlazo": obtener_texto_limpio(root, "Cnbv_DiasPlazo"),
            "AutoridadNombre": obtener_texto_limpio(root, "AutoridadNombre"),
            "NombreSolicitante": obtener_texto_limpio(root, "NombreSolicitante"),
            "Referencia": obtener_texto_limpio(root, "Referencia"),
            "Referencia1": obtener_texto_limpio(root, "Referencia1"),
            "Referencia2": obtener_texto_limpio(root, "Referencia2"),
            "TieneAseguramiento": obtener_texto_limpio(root, "TieneAseguramiento")
        },
        "SolicitudPartes": [],
        "SolicitudesEspecificas": []
    }

    for parte in root.findall('ns:SolicitudPartes', ns):
        parte_dict = {
            "ParteId": obtener_texto_limpio(parte, "ParteId"),
            "Caracter": obtener_texto_limpio(parte, "Caracter"),
            "Persona": obtener_texto_limpio(parte, "Persona"),
            "Paterno": obtener_texto_limpio(parte, "Paterno"),
            "Materno": obtener_texto_limpio(parte, "Materno"),
            "Nombre": obtener_texto_limpio(parte, "Nombre"),
            "Rfc": obtener_texto_limpio(parte, "Rfc")
        }
        datos_resultado["SolicitudPartes"].append(parte_dict)

    for sol in root.findall('ns:SolicitudEspecifica', ns):
        lista_personas_solicitud = []
        for pers in sol.findall('ns:PersonasSolicitud', ns):
            persona_dict = {
                "PersonaId": obtener_texto_limpio(pers, "PersonaId"),
                "Caracter": obtener_texto_limpio(pers, "Caracter"),
                "Persona": obtener_texto_limpio(pers, "Persona"),
                "Paterno": obtener_texto_limpio(pers, "Paterno"),
                "Materno": obtener_texto_limpio(pers, "Materno"),
                "Nombre": obtener_texto_limpio(pers, "Nombre"),
                "Rfc": obtener_texto_limpio(pers, "Rfc"),
                "Relacion": obtener_texto_limpio(pers, "Relacion"),
                "Domicilio": obtener_texto_limpio(pers, "Domicilio"),
                "Complementarios": obtener_texto_limpio(pers, "Complementarios")
            }
            lista_personas_solicitud.append(persona_dict)
        
        sol_especifica_dict = {
            "SolicitudEspecificaId": obtener_texto_limpio(sol, "SolicitudEspecificaId"),
            "InstruccionesCuentasPorConocer": obtener_texto_limpio(sol, "InstruccionesCuentasPorConocer"),
            "PersonasSolicitud": lista_personas_solicitud 
        }
        datos_resultado["SolicitudesEspecificas"].append(sol_especifica_dict)

    return datos_resultado

def extract_data_folios_to_json(folios_dict, folder_documents, json_output):

    resultados_finales = []

    for folio in folios_dict:
        
        layout_oficio = folio.get("Oficio", "")
        folio_sugo = folio.get("Folio Sugo", "")
        layout_expediente = folio.get("Expediente", "")
        layout_publicacion = folio.get("Publicacion", "")
        layout_plazo = folio.get("Plazo", "")
        layout_expediente_corto = folio.get("Expediente Corto", "")
        file_exist_caratula = folio.get("Caratula", "")
        file_exist_xml = folio.get("XML", "")
        name_file_doc = folio.get("Archivo Doc", "")
        name_file_xml = folio.get("Archivo XML", "")

        # 2. Procesar el documento de Word
        extract_name_autoridad = "No encontrado"
        if name_file_doc:
            path_file_doc = os.path.join(folder_documents, name_file_doc)
            if os.path.exists(path_file_doc):
                text_document = extractor_text_document_word(path_file_doc)
                extract_name_autoridad = search_autoridad(text_document)

        # 3. Procesar el XML o generar estructura vacía
        extract_json_xml = None
        if file_exist_xml == "OK" and name_file_xml:
            path_file_xml = os.path.join(folder_documents, name_file_xml)
            if os.path.exists(path_file_xml):
                extract_json_xml = parsear_expediente_cnbv(xml_source=path_file_xml, es_archivo=True)
            else:
                extract_json_xml = generar_estructura_vacia_xml()
        else:
            # Genera estructura vacía si no existe o el status no es "OK"
            extract_json_xml = generar_estructura_vacia_xml()

        # 4. Construir el Objeto Completo integrando todo con el prefijo "layout_"
        objeto_resultado = {
            "Oficio": layout_oficio,
            "Folio_Sugo": folio_sugo,
            "Expediente": layout_expediente,
            "Publicacion": layout_publicacion,
            "Plazo": layout_plazo,
            "Expediente_Corto": layout_expediente_corto,
            "Caratula": file_exist_caratula,
            "XML_Status": file_exist_xml,
            "Archivo_Doc": name_file_doc,
            "Archivo_XML": name_file_xml,
            
            # Datos Extraídos
            "autoridad_extraida": extract_name_autoridad,
            "xml_extraido": extract_json_xml
        }

        # Guardarlo en nuestro array principal
        resultados_finales.append(objeto_resultado)

    # 5. Exportar los datos
    
    # Guardar array de objetos a un archivo JSON (Ideal para procesar en otras aplicaciones web o APIs)
    with open(json_output, 'w', encoding='utf-8') as f:
        json.dump(resultados_finales, f, ensure_ascii=False, indent=4)


    print(f"\n[SUCCESS] Proceso terminado. ")
    print(f" -> Resultados JSON guardados en: {json_output}")

"""
    ==================================
    FUNCIONES RPA
    ==================================
"""

# --- WORKER PARA MULTIPROCESAMIENTO ---
async def procesar_lote(id_worker, oficios_lote, context, senal_inicio, pbar, lock_exito, lock_error, df, output_csv_temporal, URL_SUGO, BATCH_GUARDADO, CHECKPOINT_EVERY):
    if not oficios_lote:
        senal_inicio.set()
        return

    page = await context.new_page()

    try:
        await page.goto(URL_SUGO, wait_until="domcontentloaded", timeout=40_000)
        await asyncio.sleep(2)
        senal_inicio.set() # Avisa a la función principal que el worker ya cargó

        async def handle_dialog(dialog):
            print(f"\n[⚠️ ALERTA DETECTADA] Mensaje: {dialog.message}")
            await page.pause() 

        page.on("dialog", handle_dialog)

        for i, idx in enumerate(oficios_lote):
            oficio = df.at[idx, 'Oficio']

            folio_sugo = await get_folio_sugo(oficio=oficio, page=page) 

            df.at[idx, 'Folio Sugo'] = folio_sugo
            df.at[idx, 'Extract Folio Sugo'] = 'PROCESADO'
            pbar.update(1)

            if i > 0 and i % BATCH_GUARDADO == 0 and i % CHECKPOINT_EVERY == 0:
                async with lock_exito:
                    df.to_csv(output_csv_temporal, index=False)

    except Exception as e:
        print(f"\n[!] Error en worker {id_worker}: {e}")
        if not senal_inicio.is_set():
            senal_inicio.set()
    finally:
        async with lock_exito:
            df.to_csv(output_csv_temporal, index=False)
        await page.close()

async def enrich_layout_with_folio_sugo(excel_file, output_excel, output_csv_temporal):
    print("\n===============================================")
    print("        LOGIN EN SUGO (Para Extracción)      ")
    print("===============================================")
    user_input = input("Usuario: ").strip()
    pass_input = getpass.getpass("Contraseña: ").strip()

    if os.path.exists(output_csv_temporal):
        print("\n[INFO] Se encontro un (resultados temporal) retomando...")
        df = pd.read_csv(output_csv_temporal)
    else:
        try:
            print("\n[INFO] Leyendo nuevo excel...")
            df = pd.read_excel(excel_file, dtype=str).fillna("")
            
            columns_read = df.columns.tolist()
            missing = [col for col in COLUMS_REQUIRED if col not in columns_read]
            
            if missing:
                print(f"\n{'!'*50}")
                print(f"[ERROR] El archivo '{excel_file}' no es válido.")
                print(f"Faltan las siguientes columnas: {', '.join(missing)}")
                print(f"{'!'*50}\n")
                return
            
            df['Extract Folio Sugo'] = "SIN PROCESAR"
            df['Folio Sugo'] = ""

        except Exception as e:
            print(f"[ERROR] Leyendo el archivo Excel ({excel_file}): {e}")
            return
    
    records = df[df['Extract Folio Sugo'] == 'SIN PROCESAR'].index.tolist()

    if not records:
        print(f"[INFO] No hay registros que procesar.")
        return

    print(f"[SUCCESS] Configuración finalizada. Iniciando extraccion de {len(records)} registros")

    async with async_playwright() as p:
        browser = await p.chromium.launch(headless=True, channel="chrome")
        storage_state = await login_sugo(browser, user_input, pass_input)

        if not storage_state:
            print("[ERROR] Falló el login. Abortando proceso.")
            await browser.close()
            return

        context = await browser.new_context(
            storage_state=storage_state, ignore_https_errors=True
        )

        lotes = list(split_list(records, NUM_PESTANAS))
        tasks = []

        # Locks de escritura compartidos para todos los workers
        lock_exito = asyncio.Lock()
        lock_error = asyncio.Lock()

        bar_format = "{l_bar}{bar}| {n_fmt}/{total_fmt} [{elapsed}<{remaining}, {rate_fmt}]"
        pbar = tqdm(
            total=len(records),
            desc="Inicializando workers...",
            unit="reg",
            bar_format=bar_format,
            colour="blue",
        )

        for i, lote in enumerate(lotes):
            event_load = asyncio.Event()
            pbar.set_description(f"Cargando Worker {i+1}/{NUM_PESTANAS}")

            # Pasamos explícitamente el df y las demás variables necesarias
            task = asyncio.create_task(
                procesar_lote(
                    i + 1, lote, context, event_load, pbar, lock_exito, lock_error, 
                    df, output_csv_temporal, URL_SUGO, BATCH_GUARDADO, CHECKPOINT_EVERY
                )
            )
            tasks.append(task)
            await event_load.wait()
            
        pbar.set_description("Extrayendo Folio Sugo")
        await asyncio.gather(*tasks)
        pbar.close()
        
        df.to_excel(output_excel, index=False)
        os.remove(output_csv_temporal)
        print(f"[SUCCESS] Extracción completada. Guardado en: {output_excel}")

        return "OK"

async def login_sugo(browser, user, password):
    context = await browser.new_context(
        ignore_https_errors=True, viewport={"width": 1280, "height": 800}
    )
    page = await context.new_page()

    try:
        await page.goto(URL_LOGIN, wait_until="domcontentloaded")
        await asyncio.sleep(2)
        await page.fill(".name", user)
        await page.fill(".pass", password)
        await asyncio.sleep(1)

        async with context.expect_page() as page_info:
            if await page.locator("//p[@onclick='validaCampos()']").is_visible():
                await page.evaluate("""() => validaCampos()""")
            else:
                await page.keyboard.press("Enter")

        popup = await page_info.value
        await popup.wait_for_load_state()
        await asyncio.sleep(3)

        storage = await context.storage_state()
        await context.close()
        return storage
    except Exception as e:
        print(f"Error durante el login: {e}")
        return None

async def get_folio_sugo(oficio, page: Page):
    await page.goto(URL_ESTATUS_FOLIO, wait_until="domcontentloaded", timeout=50_000)
    await asyncio.sleep(0.3)
    
    try:
        checkbox = page.locator("#rOficio")
        await checkbox.wait_for(state="visible")

        await checkbox.focus()
        await page.keyboard.press("Space")
        await page.wait_for_selector("#fOficio.etiqueta2", state="attached", timeout=5_000)

        await page.fill("#fOficio", str(oficio))
        
        await page.evaluate("buscar();")

        await page.wait_for_selector("#panelDatos1", timeout=3000)

        folio_sugo = await page.locator("#panelDatos1 #Cons1 td:nth-child(2)").inner_text()
        
        if not folio_sugo:
            return "NO ENCONTRADO"
        
        return folio_sugo

    except PlaywrightTimeoutError:
        try:
            await page.wait_for_selector("#BTACEPTAR", timeout=40_000)
            await page.click("#BTACEPTAR")
            await asyncio.sleep(0.5)
        except Exception:
            await page.goto(URL_ESTATUS_FOLIO, wait_until="domcontentloaded")
            await asyncio.sleep(0.5)

        return "NO ENCONTRADO"

    except Exception as e:
        await page.goto(URL_ESTATUS_FOLIO, wait_until="domcontentloaded")
        await asyncio.sleep(0.5)
        
        return "NO ENCONTRADO"

"""
    ==================================
    FUNCIONES PLANTILLA
    ==================================
"""

def set_col_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

def set_row_heights_and_align(table, height):
    for row in table.rows:
        row.height = height
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM

def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def generar_documento_siti(variables: dict, dir_salida: str = "salida_documentos"):
    logo_url = "https://albertolcm.github.io/public-images/imagenes/bbva.png"
    
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Helvetica"
    style.font.size = Pt(9)

    section = doc.sections[0]
    section.top_margin = Cm(3.53)
    section.bottom_margin = Cm(1.76)
    section.left_margin = Cm(1.41)
    section.right_margin = Cm(1.41)

    # ==========================================
    # 1. INSERCIÓN DEL LOGO EN EL ENCABEZADO
    # ==========================================
    try:
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
        response = requests.get(logo_url, headers=headers, timeout=15)
        
        if response.status_code == 200:
            img_data = BytesIO(response.content)
            img_data.seek(0)
            
            header = section.header
            hdr_table = header.add_table(1, 1, Inches(5))
            hdr_table.autofit = False
            hdr_cell = hdr_table.cell(0, 0)

            tc = hdr_cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for border_name in ["top", "left", "bottom", "right"]:
                border_elm = OxmlElement(f"w:{border_name}")
                border_elm.set(qn("w:val"), "none")
                tcBorders.append(border_elm)
            tcPr.append(tcBorders)

            paragraph_header = hdr_cell.paragraphs[0]
            paragraph_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph_header.add_run().add_picture(img_data, width=Inches(1.5))
            paragraph_header.paragraph_format.space_after = Pt(0)
        else:
            print(f"Error HTTP {response.status_code} al descargar el logo.")
            doc.add_paragraph(f"<<< ERROR HTTP {response.status_code} AL OBTENER LOGO >>>")

    except Exception as e:
        print(f"Error de red o ejecución al obtener logo: {e}")
        doc.add_paragraph(f"<<< ERROR AL OBTENER LOGO: {e} >>>")

    doc.add_paragraph()

    # ==========================================
    # 2. RESTO DEL DOCUMENTO
    # ==========================================
    
    dias_semana = {
        0: "Lunes", 1: "Martes", 2: "Miércoles", 3: "Jueves",
        4: "Viernes", 5: "Sábado", 6: "Domingo"
    }

    meses = {
        1: "Enero", 2: "Febrero", 3: "Marzo", 4: "Abril",
        5: "Mayo", 6: "Junio", 7: "Julio", 8: "Agosto",
        9: "Septiembre", 10: "Octubre", 11: "Noviembre", 12: "Diciembre"
    }

    ahora = datetime.now()
    nombre_dia = dias_semana[ahora.weekday()]       
    nombre_mes = meses[ahora.month]                 

    fecha_hoy = f"{nombre_dia} {ahora.day} de {nombre_mes} de {ahora.year}"
    
    p_fecha = doc.add_paragraph(variables.get("fecha", fecha_hoy))
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph() 

    print(fecha_hoy)

    # ---------------------------------------------------------
    # ENCABEZADO Y RECUADRO DE ACUSE
    # ---------------------------------------------------------
    table_encabezado = doc.add_table(rows=1, cols=2)
    table_encabezado.autofit = False
    table_encabezado.rows[0].height = Cm(1.8)

    set_col_widths(table_encabezado, [Cm(11.48), Cm(7.252)])

    celda_izq = table_encabezado.cell(0, 0)
    p_izq = celda_izq.paragraphs[0]
    run_izq = p_izq.add_run(
        "Comisión Nacional Bancaria y de Valores\n"
        "Vicepresidencia de Supervisión de Procesos Preventivos\n"
        "Dirección General de Atención a Autoridades\n"
        'Coordinación de Atención a Autoridades "A"'
    )
    run_izq.bold = True
    run_izq.font.size = Pt(10)

    celda_der = table_encabezado.cell(0, 1)
    celda_der.vertical_alignment = WD_ALIGN_VERTICAL.CENTER 
    p_der = celda_der.paragraphs[0]
    p_der.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    run_acuse = p_der.add_run("\nAcuse de recibo CNBV")
    run_acuse.bold = True
    run_acuse.font.color.rgb = RGBColor(166, 166, 166)
    
    tc_pr = celda_der._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for border in ['top', 'left', 'bottom', 'right']:
        tag = OxmlElement(f'w:{border}')
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), '4') 
        tc_borders.append(tag)
    tc_pr.append(tc_borders)

    doc.add_paragraph() 

    p_atencion = doc.add_paragraph("Atención Dr. Jorge Alfredo Ramírez Talamantes")
    p_atencion.runs[0].bold = True

    # ---------------------------------------------------------
    # TABLA 1: Datos del Oficio
    # ---------------------------------------------------------
    table_oficio = doc.add_table(rows=4, cols=3)
    table_oficio.style = 'Table Grid'
    table_oficio.autofit = False
    
    set_col_widths(table_oficio, [Cm(2.0), Cm(4.5), Cm(12.25)])
    
    # 1. PRIMERO APLICAMOS LA ALTURA Y EL ALINEADO AL PIE A TODA LA TABLA
    set_row_heights_and_align(table_oficio, Cm(0.6))

    # 2. DESPUÉS COMBINAMOS Y CENTRAMOS LA CELDA "ASUNTO" (Para que no se sobreescriba)
    a = table_oficio.cell(0, 0)
    b = table_oficio.cell(3, 0)
    a.merge(b)
    
    a.vertical_alignment = WD_ALIGN_VERTICAL.CENTER 
    a.text = "Asunto:"
    parrafo_asunto = a.paragraphs[0]
    parrafo_asunto.runs[0].bold = True
    parrafo_asunto.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 3. FINALMENTE LLENAMOS EL RESTO DE LOS DATOS
    table_oficio.cell(0, 1).text = "Oficio:"
    table_oficio.cell(0, 2).text = variables.get("oficio", "")
    table_oficio.cell(1, 1).text = "Expediente:"
    table_oficio.cell(1, 2).text = variables.get("expediente", "")
    table_oficio.cell(2, 1).text = "Folio:"
    table_oficio.cell(2, 2).text = variables.get("folio", "")
    table_oficio.cell(3, 1).text = "Autoridad solicitante:"
    table_oficio.cell(3, 2).text = variables.get("autoridad", "")

    doc.add_paragraph()

    # ---------------------------------------------------------
    # TABLA 2: Tipos
    # ---------------------------------------------------------
    table_tipos = doc.add_table(rows=2, cols=2)
    table_tipos.style = 'Table Grid'
    table_tipos.autofit = False

    set_col_widths(table_tipos, [Cm(4.5), Cm(14.25)])
    
    table_tipos.cell(0, 0).text = "Tipo de respuesta:"
    table_tipos.cell(0, 0).paragraphs[0].runs[0].bold = True
    table_tipos.cell(0, 1).text = variables.get("tipo_respuesta", "")
    
    table_tipos.cell(1, 0).text = "Tipo de asunto:"
    table_tipos.cell(1, 0).paragraphs[0].runs[0].bold = True
    table_tipos.cell(1, 1).text = variables.get("tipo_asunto", "")

    # APLICAR ALTURA Y CENTRADO A TABLA 2
    set_row_heights_and_align(table_tipos, Cm(0.6))

    doc.add_paragraph() 

    doc.add_paragraph(
        "En atención al oficio señalado al rubro, nos permitimos hacer de su conocimiento "
        "que se ha procedido a ejecutar la instrucción de la autoridad requirente respecto "
        "a la(s) persona(s) que abajo se indica(n):"
    )

    # ---------------------------------------------------------
    # TABLA 3: Nombre y RFC
    # ---------------------------------------------------------
    table_cliente = doc.add_table(rows=1, cols=4)
    table_cliente.style = 'Table Grid'
    table_cliente.autofit = False

    set_col_widths(table_cliente, [Cm(2.5), Cm(9.5), Cm(1.5), Cm(5.25)])
    
    # 1. Aplicamos altura y alineación al pie
    set_row_heights_and_align(table_cliente, Cm(0.6))


    table_cliente.cell(0, 0).text = "NOMBRE:"
    table_cliente.cell(0, 1).text = variables.get("nombre", "")
    table_cliente.cell(0, 2).text = "RFC:"
    table_cliente.cell(0, 3).text = variables.get("rfc", "")

    color_gris = "A6A6A6" 
    
    for row in table_cliente.rows:
        for cell in row.cells:
            # Poner el fondo gris
            set_cell_background(cell, color_gris)
            # Poner todo el texto en negritas
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    

    for _ in range(1):
        doc.add_paragraph()

    # ---------------------------------------------------------
    # TABLA 4: Datos de la Cuenta
    # ---------------------------------------------------------
    table_cuentas = doc.add_table(rows=3, cols=7)
    table_cuentas.style = 'Table Grid'
    table_cuentas.autofit = False

    anchos_cuentas = [Cm(2.5), Cm(3.2), Cm(1.8), Cm(1.8), Cm(4.7), Cm(2.5), Cm(2.25)]
    set_col_widths(table_cuentas, anchos_cuentas)

    encabezados = [
        "No. Cuenta", "Tipo", "Estatus", "Carácter", 
        "Ubicación/Sucursal", "Saldo", "Moneda"
    ]
    
    color_fondo_encabezado = "D9D9D9"
    
    for i, texto in enumerate(encabezados):
        celda = table_cuentas.cell(0, i)
        celda.text = texto
        
        parrafo = celda.paragraphs[0]
        parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        parrafo.runs[0].bold = True
        
        set_cell_background(celda, color_fondo_encabezado)

    datos_cuenta = [
        variables.get("no_cuenta", ""),
        variables.get("tipo_cuenta", ""),
        variables.get("estatus_cuenta", ""),
        variables.get("caracter_cuenta", ""),
        variables.get("sucursal", ""),
        variables.get("saldo", ""),
        variables.get("moneda", "")
    ]
    
    for i, texto in enumerate(datos_cuenta):
        celda = table_cuentas.cell(1, i)
        celda.text = texto
        celda.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    celda_obs_titulo = table_cuentas.cell(2, 0)
    celda_obs_titulo.text = "Observaciones"
    celda_obs_titulo.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    celda_obs_texto = table_cuentas.cell(2, 1)
    celda_obs_texto.merge(table_cuentas.cell(2, 6))
    
    texto_obs = variables.get("observaciones", "")
    celda_obs_texto.text = texto_obs
    celda_obs_texto.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    for row in table_cuentas.rows:
        row.height = Cm(0.8) 
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()


    for _ in range(1):
        doc.add_paragraph()


    # ==========================================
    # BLOQUE DE FIRMA
    # ==========================================
    p_linea = doc.add_paragraph("______________________________________")
    p_linea.alignment = WD_ALIGN_PARAGRAPH.CENTER

    datos_firma = [
        "VANESSA RAMIREZ PATIÑO",
        "GERENTE",
        "BBVA MÉXICO, S.A.",
        "INSTITUCIÓN DE BANCA MÚLTIPLE",
        "GRUPO FINANCIERO BBVA MÉXICO"
    ]

    for line in datos_firma:
        p_data = doc.add_paragraph(line)
        p_data.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_data.paragraph_format.space_before = Pt(0)
        p_data.paragraph_format.space_after = Pt(0)

    doc.add_paragraph()
    p_dof = doc.add_paragraph(f"{variables.get('variable_dof', '')}-AJRT")
    p_dof.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Guardado del documento
    try:
        os.makedirs(dir_salida, exist_ok=True)
        folio_seguro = variables.get('folio', 'SIN_FOLIO').replace("/", "-")
        expediente_corto = variables.get('expediente_corto')
        nombre_archivo = f"{expediente_corto}.docx"
        ruta_completa = os.path.join(dir_salida, nombre_archivo)
        
        doc.save(ruta_completa)
        print(f"Archivo generado exitosamente en: {ruta_completa}")
        return ruta_completa
        
    except Exception as e:
        print(f"Error crítico al intentar guardar el documento: {e}")
        return None


"""
    ==================================
    FUNCIONES DOWNLOAD FILES CNBV
    ==================================
"""

def cargar_credenciales():
    """Carga y devuelve las credenciales del archivo JSON."""
    try:
        with open(ARCHIVO_CREDENCIALES, 'r') as f:
            credenciales = json.load(f)
        return credenciales['usuario'], credenciales['contrasena']
    except FileNotFoundError:
        print(f"❌ Error: No se encontró el archivo '{ARCHIVO_CREDENCIALES}'.")
        sys.exit(1)
    except KeyError as e:
        print(f"❌ Error: Formato de credenciales inválido. Falta la clave {e}.")
        sys.exit(1)
    except Exception as e:
        print(f"❌ Error al cargar credenciales: {e}")
        sys.exit(1)

def preparar_entorno():
    """Limpia resultados anteriores y crea el directorio de descargas."""
    if os.path.exists(ARCHIVO_RESULTADOS):
        os.remove(ARCHIVO_RESULTADOS)
    if os.path.exists(FOLDER_DOCUMENTS):
        shutil.rmtree(FOLDER_DOCUMENTS)
    os.makedirs(FOLDER_DOCUMENTS)
    print("✅ Entorno preparado.")

def obtener_fechas_terminal():
    """Solicita y valida las fechas desde la terminal."""
    print("\n--- Configuración de Fechas ---")
    print("Formato requerido: AAAA-MM-DD (ejemplo: 2023-01-31)\n")
    
    def pedir_fecha(mensaje):
        while True:
            fecha_str = input(mensaje).strip()
            try:
                fecha_obj = datetime.strptime(fecha_str, '%Y-%m-%d')
                return fecha_obj, fecha_str
            except ValueError:
                print("❌ Formato incorrecto. Por favor usa el formato AAAA-MM-DD.")
                
    fecha_inicio_obj, fecha_inicio_str = pedir_fecha("Ingresa la Fecha de INICIO: ")
    
    while True:
        fecha_fin_obj, fecha_fin_str = pedir_fecha("Ingresa la Fecha de FIN: ")
        if fecha_fin_obj >= fecha_inicio_obj:
            break
        print("❌ La fecha de fin debe ser igual o posterior a la fecha de inicio.")
        
    return fecha_inicio_str, fecha_fin_str

def intentar_descarga(page, locator, ruta_destino):
    """
    Intenta realizar una descarga usando Playwright interceptando el evento.
    Devuelve "OK" si es exitoso, "ERROR" si falla.
    """
    try:
        # Verifica si el enlace realmente existe antes de intentar hacer clic
        if locator.count() == 0:
            return "NO DISPONIBLE"
            
        with page.expect_download(timeout=15000) as download_info:
            locator.click()
            
        download = download_info.value
        # Guarda el archivo en la ruta específica
        download.save_as(ruta_destino)
        return "OK"
    except PlaywrightTimeoutError:
        return "ERROR (Tiempo de espera)"
    except Exception as e:
        return f"ERROR ({str(e)})"

def ejecutar_automatizacion(usuario, password, fecha_inicio, fecha_fin):
    """Lógica principal de navegación y extracción de datos usando Playwright."""
    resultados = []
    
    # Iniciamos Playwright
    with sync_playwright() as p:
        # headless=False permite ver el navegador. Cambia a True si quieres que corra en segundo plano.
        browser = p.chromium.launch(headless=False) 
        context = browser.new_context(accept_downloads=True)
        page = context.new_page()

        try:
            print("\n🚀 Iniciando navegador e iniciando sesión...")
            page.goto("https://websitiaa.cnbv.gob.mx/logOn.aspx")
            page.fill("#ctl00_DefaultPlaceholder_textBoxUser", usuario)
            page.fill("#ctl00_DefaultPlaceholder_textBoxPassword", password)
            page.click("#ButtonValidate")
            
            # Esperamos a que la navegación post-login termine
            page.wait_for_load_state("networkidle")

            print(f"🔎 Buscando registros del {fecha_inicio} al {fecha_fin}...")
            page.goto("https://websitiaa.cnbv.gob.mx/Publicados.aspx")
            
            # Llenamos el formulario de búsqueda
            page.select_option("#ctl00_DefaultPlaceholder_ComboBoxAreas", label="Aseguramiento")
            page.select_option("#ctl00_DefaultPlaceholder_ComboBoxEstatusOficio", label="Todos")
            page.fill("#ctl00_DefaultPlaceholder_TextFechaPublicacion1", fecha_inicio)
            page.fill("#ctl00_DefaultPlaceholder_TextFechaPublicacion2", fecha_fin)
            page.click("#ctl00_DefaultPlaceholder_ButtonQuery")

            # Esperamos a que aparezca la tabla de resultados
            try:
                page.wait_for_selector("#ctl00_DefaultPlaceholder_GridResult", timeout=10000)
            except PlaywrightTimeoutError:
                print("⚠️ No se encontró la tabla de resultados (posiblemente no hay registros en esas fechas).")
                return resultados

            # Obtenemos todas las filas de la tabla
            filas = page.locator("#ctl00_DefaultPlaceholder_GridResult tr").all()
            total_oficios = len(filas) - 1 # Restamos el encabezado
            
            print(f"📊 Total de oficios encontrados: {total_oficios}")

            # Procesamos cada fila (saltamos el índice 0 que es el encabezado)
            for i, fila in enumerate(filas[1:], start=1):
                columnas = fila.locator("td").all()
                
                # Aseguramos que la fila tenga la cantidad esperada de columnas
                if len(columnas) < 10:
                    continue
                
                # Extraemos el texto de las columnas relevantes
                oficio = columnas[5].inner_text().strip()
                expediente = columnas[6].inner_text().strip()
                publicacion = columnas[7].inner_text().strip()
                plazo = columnas[9].inner_text().strip()

                expediente_corto = expediente[-10:] if len(expediente) >= 10 else expediente
                
                # Nombres de archivos limpios para guardarlos
                archivo_doc_nombre = f"{oficio}.doc".replace("/", "-")
                archivo_xml_nombre = f"{oficio}.xml".replace("/", "-")
                
                ruta_doc = os.path.join(FOLDER_DOCUMENTS, archivo_doc_nombre)
                ruta_xml = os.path.join(FOLDER_DOCUMENTS, archivo_xml_nombre)

                print(f"[{i}/{total_oficios}] Procesando Oficio: {oficio}")

                # Enlaces de descarga
                enlace_caratula = columnas[-3].locator("a")
                enlace_xml = columnas[-1].locator("a")

                # Descargamos los archivos usando nuestra función robusta
                estado_caratula = intentar_descarga(page, enlace_caratula, ruta_doc)
                estado_xml = intentar_descarga(page, enlace_xml, ruta_xml)

                print(f"    ↳ Carátula: {estado_caratula} | XML: {estado_xml}")

                # Guardamos el registro en la lista de resultados
                resultados.append({
                    "Oficio": oficio,
                    "Expediente": expediente,
                    "Publicacion": publicacion,
                    "Plazo": plazo,
                    "Expediente Corto": expediente_corto,
                    "Caratula": estado_caratula,
                    "XML": estado_xml,
                    "Archivo Doc": archivo_doc_nombre if estado_caratula == "OK" else "",
                    "Archivo XML": archivo_xml_nombre if estado_xml == "OK" else ""
                })

        except Exception as e:
            print(f"\n❌ Ocurrió un error crítico durante la automatización: {e}")
        finally:
            context.close()
            browser.close()
            
    return resultados

def download_files_cnbv():
    print("=========================================")
    print("   Bot de Descarga de Acuses (CNBV)      ")
    print("=========================================\n")
    
    # 1. Cargar credenciales
    usuario, password = cargar_credenciales()
    
    # 2. Solicitar fechas en terminal
    fecha_inicio, fecha_fin = obtener_fechas_terminal()
    
    # 3. Limpiar y preparar carpetas
    preparar_entorno()
    
    # 4. Ejecutar la extracción con Playwright
    datos_extraidos = ejecutar_automatizacion(usuario, password, fecha_inicio, fecha_fin)
    
    # 5. Generar archivo Excel si hay resultados
    if datos_extraidos:
        df_resultados = pd.DataFrame(datos_extraidos)
        df_resultados.to_excel(ARCHIVO_RESULTADOS, index=False)
        print(f"\n✅ Proceso completado con éxito. Se procesaron {len(datos_extraidos)} registros.")
        print(f"📁 Los resultados han sido guardados en '{ARCHIVO_RESULTADOS}'.")
        print(f"📁 Los archivos descargados están en la carpeta '{FOLDER_DOCUMENTS}'.")
    else:
        print("\n⚠️ El proceso finalizó sin resultados para exportar.")

if __name__ == "__main__":
    print_banner_ascii()

    download_files_cnbv()

    if os.path.exists(FOLDER_GENERATES):
        shutil.rmtree(FOLDER_GENERATES)
        print(f"[INFO] La carpeta '{FOLDER_GENERATES}' y sus archivos fueron eliminados.")

    # 2. Volver a crear la carpeta (vacía)
    os.makedirs(FOLDER_GENERATES)
    print(f"[SUCCESS] Carpeta '{FOLDER_GENERATES}' creada nuevamente y lista para usarse.")

    resultado = asyncio.run(enrich_layout_with_folio_sugo(INPUT_FILE, OUTPUT_EXCEL, OUTPUT_TEMPORAL))

    if not resultado == "OK":
        print("[ERROR] Ocurrio algo al extraer SUGO. Intenta de nuevo...")
        sys.exit(1)
    
    # Comprobación del Lay Out (revisar si tiene el FOLIO SUGO)
    try:
        print("\n[INFO] Validando Layout...")

        df = pd.read_excel(OUTPUT_EXCEL, dtype=str).fillna("")
        
        columns_read = df.columns.tolist()
        missing = [col for col in COLUMNS_WITH_SUGO if col not in columns_read]
        
        if missing:
            print(f"\n{'!'*50}")
            print(f"[ERROR] El archivo '{OUTPUT_EXCEL}' no es válido.")
            print(f"Faltan las siguientes columnas: {', '.join(missing)}")
            print(f"{'!'*50}\n")
            sys.exit(1)

    except Exception as e:
        print(f"[ERROR] Leyendo el archivo Excel ({OUTPUT_EXCEL}): {e}")
        sys.exit(1)
    
    empty_values = df['Folio Sugo'].isna() | (df['Folio Sugo'].astype(str).str.strip() == "")

    if empty_values.any():
        filas_con_error = df[empty_values].index.tolist()
        print(f"\n[ERROR] Se detectaron {len(filas_con_error)} registros con el 'Folio Sugo' vacío.")
        print(f"Filas afectadas: {filas_con_error}")
        
        raise SystemExit("Ejecución detenida: Faltan datos en la columna 'Folio Sugo'.")

    print("[SUCCESS] Validación superada. Todos los registros tienen Folio Sugo.") 

    print("\n[INFO] Extrayendo informacion...")
    folios_dict = df.to_dict('records')
    extract_data_folios_to_json(
        folder_documents=FOLDER_DOCUMENTS, 
        folios_dict=folios_dict, 
        json_output=OUTPUT_JSON
    )




    with open(OUTPUT_JSON, 'r', encoding='utf-8') as archivo:
        data = json.load(archivo)

    for oficio in data:
        # 1. Definimos valores por defecto en caso de que vengan vacíos
        nombre_extraido = "NO ENCONTRADO"
        rfc_extraido = "SIN RFC"
        
        # 2. Navegación segura por el JSON
        xml = oficio.get("xml_extraido", {})
        solicitudes = xml.get("SolicitudesEspecificas", [])
        
        # 3. Verificamos que la lista 'solicitudes' exista y tenga al menos 1 elemento
        if solicitudes and len(solicitudes) > 0:
            personas = solicitudes[0].get("PersonasSolicitud", [])
            
            # Verificamos que la lista 'personas' exista y tenga al menos 1 elemento
            if personas and len(personas) > 0:
                nombre_extraido = personas[0].get("Nombre", "NO ENCONTRADO")
                rfc_extraido = personas[0].get("Rfc", "SIN RFC")

        # 4. Armamos nuestro diccionario final usando las variables seguras
        mis_variables = {
            "oficio": oficio.get("Oficio"),
            "expediente": "",
            "folio": oficio.get("Expediente"),
            "expediente_corto": oficio.get("Expediente_Corto"),
            "autoridad": oficio.get("autoridad_extraida"),
            "tipo_respuesta": "Total",
            "tipo_asunto": "Aseguramiento",
            "nombre": nombre_extraido,     # <- Variable segura
            "rfc": rfc_extraido,           # <- Variable segura
            "variable_dof": oficio.get("Folio_Sugo") 
        }

        print("[LOG] Word Generado")
        generar_documento_siti(mis_variables)